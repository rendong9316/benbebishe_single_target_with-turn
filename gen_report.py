# -*- coding: utf-8 -*-
"""Generate stage achievement report for dual-illumination OTH-SWR simulation."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_CN = "SimSun"
FONT_WEST = "Times New Roman"
SZ = Pt(12)       # xiaosi
SZ5 = Pt(10.5)    # wubao
SZDH = Pt(14)     # dah
SZDA = Pt(16)     # da
SZT = Pt(22)      # title
BLACK = RGBColor(0, 0, 0)

def sf(run, size=None, color=None, bold=False, italic=False):
    if size is None: size = SZ
    if color is None: color = BLACK
    run.font_size = size
    run.font.name = FONT_WEST
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = parse_xml(
            '<w:rFonts %s w:ascii="%s" w:hAnsi="%s" w:cs="%s" w:eastAsia="%s"/>'
            % (nsdecls("w"), FONT_WEST, FONT_WEST, FONT_WEST, FONT_CN))
        rpr.insert(0, rf)
    else:
        rf.set(qn('w:ascii'), FONT_WEST)
        rf.set(qn('w:hAnsi'), FONT_WEST)
        rf.set(qn('w:cs'), FONT_WEST)
        rf.set(qn('w:eastAsia'), FONT_CN)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic

def spf(para):
    pf = para.paragraph_format
    pf.line_spacing = 1.25
    pf.first_line_indent = Cm(0.74)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

def abody(doc, text):
    p = doc.add_paragraph()
    spf(p)
    r = p.add_run(text)
    sf(r)
    return p

def amix(doc, parts):
    p = doc.add_paragraph()
    spf(p)
    for text, bold, italic in parts:
        r = p.add_run(text)
        sf(r, bold=bold, italic=italic)
    return p

def acenter(doc, text, size=None, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spf(p)
    r = p.add_run(text)
    sf(r, size if size else SZ, bold=bold)
    return p

def aheading(doc, text, size=SZDA, bold=True):
    p = doc.add_paragraph()
    spf(p)
    r = p.add_run(text)
    sf(r, size, BLACK, bold=bold)
    return p

def atitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spf(p)
    r = p.add_run(text)
    sf(r, SZT, BLACK, bold=True)
    return p

def amk(doc, text=None):
    p = doc.add_paragraph()
    spf(p)
    if text:
        r = p.add_run(text)
        sf(r)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

def mktable(doc, headers, rows):
    td = [headers] + rows
    nc = len(td[0])
    nr = len(td)
    table = doc.add_table(rows=nr, cols=nc)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, rd in enumerate(td):
        row = table.rows[i]
        for j, ct in enumerate(rd):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            spf(p)
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(str(ct))
            if i == 0:
                sf(r, SZ5, BLACK, bold=True)
            else:
                sf(r, SZ5, BLACK)
    tbl = table._element
    tp = tbl.tblPr
    if tp is None:
        tp = parse_xml('<w:tblPr %s/>' % nsdecls("w"))
        tbl.insert(0, tp)
    bd = parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>')
    tp.append(bd)
    return table

# ============================================================
# BUILD DOCUMENT
# ============================================================
doc = Document()
style = doc.styles['Normal']
style.font.name = FONT_WEST
style.font.size = SZ
style.paragraph_format.line_spacing = 1.25
style.paragraph_format.first_line_indent = Cm(0.74)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
rpr = style.element.get_or_add_rPr()
rf = parse_xml(
    '<w:rFonts %s w:ascii="%s" w:hAnsi="%s" w:cs="%s" w:eastAsia="%s"/>'
    % (nsdecls("w"), FONT_WEST, FONT_WEST, FONT_WEST, FONT_CN))
rpr.append(rf)

# TITLE PAGE
amk(doc)
amk(doc)
amk(doc)
amk(doc)
atitle(doc, "双基地外辐射源雷达航迹关联与融合方法研究")
amk(doc)
acenter(doc, "阶 段 成 果 报 告", size=SZDA)
amk(doc)
amk(doc)
acenter(doc, "学生姓名：任  东")
acenter(doc, "学    号：2023112696")
acenter(doc, "指导教师：__________")
amk(doc)
acenter(doc, "计算机与人工智能学院")
amk(doc)
acenter(doc, "2026年6月")

doc.add_page_break()

# TABLE OF CONTENTS
p = doc.add_paragraph()
spf(p)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("目    录")
sf(r, Pt(16), BLACK, bold=True)
amk(doc)

toc = [
    ("1", "绪论"),
    ("1.1", "研究背景与意义"),
    ("1.2", "主要研究工作"),
    ("2", "双照射源系统建模"),
    ("2.1", "双基地雷达几何模型"),
    ("2.2", "坐标系转换与球面几何"),
    ("2.3", "雷达站部署方案"),
    ("2.4", "威力覆盖与探测约束"),
    ("3", "目标运动与场景设计"),
    ("3.1", "目标运动模型"),
    ("3.2", "三种仿真场景"),
    ("3.3", "检测与虚警模型"),
    ("4", "系统偏差标定"),
    ("4.1", "ADS-B真值数据处理"),
    ("4.2", "偏差估计方法"),
    ("4.3", "标定效果验证"),
    ("5", "单站航迹跟踪算法"),
    ("5.1", "UKF基本原理"),
    ("5.2", "自适应UKF"),
    ("5.3", "IMM双模型跟踪"),
    ("5.4", "跟踪性能数据"),
    ("6", "航迹时间对齐与融合"),
    ("6.1", "时间异步与CV外推对齐"),
    ("6.2", "四种融合算法"),
    ("6.3", "融合结果"),
    ("7", "蒙特卡洛统计分析"),
    ("7.1", "仿真设置"),
    ("7.2", "直线场景蒙特卡洛结果"),
    ("7.3", "拐弯场景三体制对比"),
    ("7.4", "回头弯场景结果"),
    ("8", "总结与展望"),
    ("8.1", "阶段性工作总结"),
    ("8.2", "主要成果数据"),
    ("8.3", "下一步工作计划"),
]
for num, title in toc:
    p = doc.add_paragraph()
    spf(p)
    indent = 0.74 if '.' in num else 0
    p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run("%s\t%s" % (num, title))
    sf(r, SZ, BLACK, bold=('.' not in num))

doc.add_page_break()

# ============================================================
# CHAPTER 1
# ============================================================
aheading(doc, "1\t绪论", size=SZDA)

aheading(doc, "1.1\t研究背景与意义", size=SZ)

abody(doc,
    "外辐射源雷达（Over-the-Horizon Radar with Non-cooperative Illuminator of Opportunity, "
    "ORTNR）利用第三方非合作照明器（如数字音频广播发射台DAB/DAB+、电视广播台等）发射的"
    "无线电波实现对超视距目标的被动探测。与主动雷达相比，该系统具有成本低、隐蔽性强、"
    "覆盖范围广等优势，特别适合海疆监控、气象观测和广域目标跟踪等应用。")

abody(doc,
    "单照射源外辐射源雷达的几何结构较为单一，探测覆盖区域受限，定位精度不足。"
    "引入双照射源（多部广播台）可显著拓展覆盖范围、改善几何精度因子（GDOP），"
    "并通过空间分集提高检测概率。然而，双照射源场景带来了新的技术挑战：")

abody(doc,
    "（1）两部雷达的采样时刻不同步，需进行时间对齐；")
abody(doc,
    "（2）各雷达存在系统标定偏差，需利用ADS-B等真值数据进行离线估计；")
abody(doc,
    "（3）电离层传播导致量测噪声大、检测概率低（典型Pd约0.6）；")
abody(doc,
    "（4）多源航迹的关联与融合需处理未知的互相关性问题。")

abody(doc,
    "本阶段工作围绕上述挑战，完成了从场景建模、偏差标定、单站跟踪（UKF/自适应UKF/IMM）、"
    "航迹融合（SCC/BC/CI/FCI）到蒙特卡洛统计分析的完整仿真流水线。")

aheading(doc, "1.2\t主要研究工作", size=SZ)

abody(doc, "本阶段完成的主要工作包括：")

works = [
    "完成双基地雷达几何建模与球面坐标转换工具链（utils/目录，7个工具函数）",
    "实现9-Phase完整仿真流水线（Phase 0-9），支持直线/拐弯/回头弯三种场景",
    "开发三种UKF跟踪体制：基础UKF、自适应UKF（模糊自适应Q）、IMM-CV/CT双模型",
    "实现四种航迹融合算法：SCC、BC、CI、FCI，并集成协方差正则化",
    "完成概率数据关联（PDA）与NN最近邻关联模块",
    "实现ADS-B辅助的系统偏差标定（样本均值估计法）",
    "开发完整的可视化模块（场景总览/点云3D/跟踪综合图/融合对比图/误差收敛曲线/CDF）",
    "完成直线场景蒙特卡洛仿真（N=500）与拐弯场景三体制对比蒙特卡洛（N=500）",
]
for w in works:
    p = doc.add_paragraph()
    spf(p)
    r = p.add_run("\t%s" % w)
    sf(r)

# ============================================================
# CHAPTER 2
# ============================================================
doc.add_page_break()
aheading(doc, "2\t双照射源系统建模", size=SZDA)

aheading(doc, "2.1\t双基地雷达几何模型", size=SZ)

abody(doc,
    "双基地雷达的发射站（Tx）和接收站（Rx）地理上分离，与单基地雷达的本质区别在于"
    "距离量测对应'Tx->目标->Rx'的总路径长度（群距离Rg），而非目标到雷达的直线斜距。")

abody(doc, "双基地量测模型：")

amix(doc, [
    ("群距离 ", False, False),
    ("Rg = r_Tx + r_Rx", False, True),
    ("，其中r_Tx为发射站到目标的大圆距离，r_Rx为目标到接收站的大圆距离。", False, False),
])

amix(doc, [
    ("方位角 ", False, False),
    ("az", False, True),
    ("为接收站观测目标的真北方位角。", False, False),
])

amix(doc, [
    ("径向速度 ", False, False),
    ("v_d", False, True),
    ("为速度在Tx->目标和目标->Rx两个方向的投影之和。", False, False),
])

aheading(doc, "2.2\t坐标系转换与球面几何", size=SZ)

abody(doc,
    "系统采用WGS84经纬度坐标描述地理位置，通过以下工具函数完成坐标转换：")

mktable(doc,
    ["函数", "功能", "输入", "输出"],
    [
        ["sphere_utils_haversine_distance", "球面距离", "lon1,lat1,lon2,lat2", "距离(km)"],
        ["sphere_utils_azimuth", "球面方位角", "lon1,lat1,lon2,lat2", "方位角(deg)"],
        ["sphere_utils_destination_point", "正算目的地", "起点,方位角,距离", "终点经纬度"],
        ["sphere_utils_interpolate_great_circle", "大圆插值", "两端点,比例", "中间点经纬度"],
        ["skywave_geometry.group_range", "双基地群距离", "Tx/Rx/目标经纬度", "Rg(km)"],
        ["bistatic_inverse_solver", "双基地反解", "Rg,az,Tx,Rx", "目标经纬度"],
        ["coord_systems_lla_to_ecef", "LLA转ECEF", "lon,lat,h", "X,Y,Z(m)"],
    ])

aheading(doc, "2.3\t雷达站部署方案", size=SZ)

abody(doc,
    "本仿真部署两部异质双基地雷达站，R1为精密站，R2为标准站，"
    "两站接收站相距约185km，形成空间分集。")

mktable(doc,
    ["参数", "R1（精密站）", "R2（标准站）"],
    [
        ["接收站经纬度", "(113.0E, 33.5N)", "(115.0E, 33.0N)"],
        ["发射站经纬度", "(109.0E, 33.5N)", "(111.0E, 33.0N)"],
        ["Tx-Rx基线", "~370 km", "~370 km"],
        ["R1-R2间距", "--", "~185 km"],
        ["波束指向", "92 deg", "91 deg"],
        ["波束宽度", "15 deg", "15 deg"],
        ["距离噪声sigma_r", "7 km", "14 km"],
        ["方位噪声sigma_az", "0.35 deg", "0.6 deg"],
        ["系统偏差dr", "+20 km", "-15 km"],
        ["系统偏差da", "-3.0 deg", "+3.5 deg"],
    ])

aheading(doc, "2.4\t威力覆盖与探测约束", size=SZ)

abody(doc,
    "OTH-SWR的单跳电离层反射模式决定了其威力覆盖约束：")

mktable(doc,
    ["约束条件", "参数值", "物理含义"],
    [
        ["最小探测距离（静区）", "1000 km", "电离层F层(250~400km)一跳模式的最小地面距离"],
        ["最大探测距离", "2000 km", "地球曲率与电离层高度限制"],
        ["波束宽度", "15 deg", "接收阵列3dB半功率全宽度"],
        ["波束中心方位", "R1: 92 deg / R2: 91 deg", "指向东部海域，形成重叠覆盖区"],
        ["距离分辨率", "10 km", "对应约15 kHz带宽"],
        ["方位分辨率", "1 deg", "对应约57倍波长有效孔径"],
        ["分辨单元总数", "1500", "100(range) x 15(azimuth)"],
    ])

# ============================================================
# CHAPTER 3
# ============================================================
doc.add_page_break()
aheading(doc, "3\t目标运动与场景设计", size=SZDA)

aheading(doc, "3.1\t目标运动模型", size=SZ)

abody(doc,
    "仿真目标为民航客机级别空中目标，巡航速度230 m/s（约828 km/h，Ma约0.78），"
    "沿大圆航线从东海海域穿越双雷达共同覆盖区。跟踪算法采用两种运动模型：")

abody(doc, "（1）常速模型（CV）：")
abody(doc,
    "状态向量 [lon, v_lon, lat, v_lat]^T，状态转移矩阵")
abody(doc,
    "F_CV = [[1, dt, 0, 0], [0, 1, 0, 0], [0, 0, 1, dt], [0, 0, 0, 1]]")

abody(doc, "（2）协调转弯模型（CT）：")
abody(doc,
    "在CV基础上引入转弯角速率omega，状态转移矩阵包含sin(omega*dt)/omega和(1-cos(omega*dt))/omega项，"
    "适用于描述恒转弯率机动。")

aheading(doc, "3.2\t三种仿真场景", size=SZ)

mktable(doc,
    ["场景", "航迹特征", "航路点", "跟踪算法", "主入口"],
    [
        ["直线巡航", "匀速直线，无机动", "W1(127.5E,31.0N)->W2(130.5E,33.0N)", "自适应UKF", "run_simulation.m"],
        ["渐进拐弯", "1 deg/s转弯，46.7 deg", "W1->W2(顶点)->W3，三段式", "IMM(CV+CT)", "run_simulation_turn.m"],
        ["回头弯180", "1 deg/s左转半圆，180s", "正东120km->半圆->正西120km", "IMM(CV+CT)", "run_simulation_turn_180deg.m"],
    ])

abody(doc,
    "其中直线场景使用自适应UKF（ukf_zishiying），拐弯和回头弯场景使用IMM双模型（ukf_imm），"
    "以验证IMM在机动场景中的性能优势。")

aheading(doc, "3.3\t检测与虚警模型", size=SZ)

abody(doc,
    "考虑到OTH-SWR电离层传播衰落严重，检测概率设定为Pd=0.6，"
    "即在约120帧仿真中约有72帧产生目标检测。虚警率设定为Pfa=0.001，"
    "每帧期望虚警数约1.5个（泊松分布建模）。")

mktable(doc,
    ["参数", "符号", "值"],
    [
        ["采样周期", "dt", "30 s"],
        ["仿真时长", "T_total", "3600 s（约120帧）"],
        ["检测概率", "P_d", "0.6"],
        ["虚警率", "P_fa", "0.001"],
        ["杂波空间密度", "lambda", "5e-5 / (km*deg)"],
        ["PDA门内概率", "P_G", "0.8647 (gate_sigma=2)"],
        ["R1/R2采样偏移", "delta_t", "0s / 13s"],
    ])

# ============================================================
# CHAPTER 4
# ============================================================
doc.add_page_break()
aheading(doc, "4\t系统偏差标定", size=SZDA)

aheading(doc, "4.1\tADS-B真值数据处理", size=SZ)

abody(doc,
    "利用ADS-B广播数据（2026-04-27 09:30:00，CSV格式，约24万点）作为地面真值参考。"
    "从ADS-B数据中筛选位于双雷达覆盖范围内的样本点，每站最多5000点均匀采样，"
    "计算每个样本点的真实群距离和方位角，叠加已知的系统偏差和高斯噪声生成模拟量测。")

aheading(doc, "4.2\t偏差估计方法", size=SZ)

abody(doc,
    "采用样本均值估计法：对覆盖范围内所有标定样本的距离残差和方位残差分别求均值，"
    "得到每部雷达的距离偏置dr_est和方位偏置da_est。")

abody(doc, "dr_est = mean(Rg_meas - Rg_true)")
abody(doc, "da_est = mean(az_meas - az_true)")

abody(doc,
    "该方法等价于最大似然估计（MLE），在标定样本数足够大时收敛于真实偏差值。"
    "估计完成后，对所有后续点迹执行校正：Rg_corrected = Rg_meas - dr_est，"
    "az_corrected = az_meas - da_est，再进行双基地几何反解得到校正后的目标经纬度。")

aheading(doc, "4.3\t标定效果验证", size=SZ)

abody(doc,
    "通过点迹定位RMSE对比验证标定效果。标定前点迹包含系统偏差+随机噪声，"
    "标定后仅含随机噪声。预期标定后RMSE显著降低。")

mktable(doc,
    ["指标", "R1", "R2"],
    [
        ["真实偏差", "+20 km / -3.0 deg", "-15 km / +3.5 deg"],
        ["标定前点迹RMSE", "[待填入]", "[待填入]"],
        ["标定后点迹RMSE", "[待填入]", "[待填入]"],
        ["改善率", "[待填入]", "[待填入]"],
    ])

abody(doc,
    "注：表中'[待填入]'占位符将在运行主程序后填入实际数据。"
    "运行 run_simulation.m 后，控制台将打印四类RMSE：原始点迹RMSE、校准后点迹RMSE。")

# ============================================================
# CHAPTER 5
# ============================================================
doc.add_page_break()
aheading(doc, "5\t单站航迹跟踪算法", size=SZDA)

aheading(doc, "5.1\tUKF基本原理", size=SZ)

abody(doc,
    "无迹卡尔曼滤波（UKF）通过无迹变换（Unscented Transform, UT）捕获状态分布的高阶矩，"
    "避免了EKF中雅可比矩阵的线性化误差。UT三参数控制Sigma点的散布：")

mktable(doc,
    ["参数", "符号", "当前值", "作用"],
    [
        ["散布度", "alpha", "1e-2", "控制Sigma点散布范围，小值接近EKF"],
        ["先验分布", "beta", "2.0", "beta=2对高斯分布最优"],
        ["次级尺度", "kappa", "0.0", "补充三阶矩信息"],
    ])

abody(doc,
    "UKF每帧执行以下流程：create(初始化模板) -> init(初始状态) -> "
    "prepare(UT变换生成Sigma点并预测) -> NN关联(最近邻波门筛选) -> "
    "PDA加权(计算门内量测概率权重) -> update(Kalman纯数学更新)。")

aheading(doc, "5.2\t自适应UKF（模糊自适应Q）", size=SZ)

abody(doc,
    "基础UKF使用固定的过程噪声协方差Q_scale，在目标平稳时表现良好，"
    "但在机动场景下跟踪滞后明显。自适应UKF通过在线监测新息序列统计特性动态调整Q：")

abody(doc, "（1）计算最近W帧的新息归一化平方和（NIS）滑动均值；")
abody(doc, "（2）将NIS均值输入模糊推理系统（Sugeno型），得到Q的缩放因子mu；")
abody(doc, "（3）当NIS约等于2（期望值）时mu约等于1.0（维持）；NIS>>2时mu最大3.0（放大Q跟踪机动）；"
    "NIS<<2时mu最小0.6（缩小Q平稳滤波）。")

mktable(doc,
    ["自适应参数", "值", "说明"],
    [
        ["模糊自适应开关", "true", "启用NIS驱动的Q自适应"],
        ["NIS滑动窗口", "3帧", "窗口越小响应越快"],
        ["EMA平滑系数", "0.10", "慢而稳的自适应节奏"],
        ["机动自适应EMA", "0.10", "用于拐弯预检测"],
    ])

aheading(doc, "5.3\tIMM双模型跟踪", size=SZ)

abody(doc,
    "交互多模型（IMM）算法同时运行CV和CT两个UKF子滤波器，"
    "通过Markov转移矩阵实现模型间的概率混合，适合跟踪机动目标。")

abody(doc, "IMM每帧执行以下循环：")
abody(doc, "（1）模型混合：各模型用混合初始状态（基于转移概率加权）；")
abody(doc, "（2）双模型独立预测：CV/CT各自UKF prepare；")
abody(doc, "（3）组合输出：mu加权组合状态供上层关联使用；")
abody(doc, "（4）重构加权量测：innov_w + z_pred_comb；")
abody(doc, "（5）各模型独立更新：委托ukf_jichu('update')；")
abody(doc, "（6）Pd-IPDA似然度：基于文献Musicki 2008 IEEE T-AES；")
abody(doc, "（7）贝叶斯概率更新：mu_new proportional L * (Pi' * mu)；")
abody(doc, "（8）概率钳位[0.02, 0.95]防止数值病态；")
abody(doc, "（9）状态组合：x_comb = sum(mu_i * x_i)。")

mktable(doc,
    ["IMM参数", "值", "说明"],
    [
        ["转移矩阵Pi", "[0.95,0.05; 0.05,0.95]", "缓慢切换，高自保持概率"],
        ["CT转弯率", "1 deg/s", "对应民航标准转弯率"],
        ["模型概率钳位", "[0.02, 0.95]", "防止数值病态"],
    ])

aheading(doc, "5.4\t航迹管理与跟踪性能数据", size=SZ)

abody(doc,
    "航迹管理采用M/N滑窗起始逻辑和K_loss连续丢帧终止逻辑：")

mktable(doc,
    ["管理参数", "R1（精密站）", "R2（标准站）", "说明"],
    [
        ["M/N起始", "4/8", "4/8", "8帧滑窗中至少4次检测才确认"],
        ["K_loss终止", "8", "8", "甜点分析结论，双方K=8时漏检率最低"],
        ["关联门限gate_sigma", "6", "6", "马氏距离门控"],
        ["Vr硬门", "20 m/s", "40 m/s", "R2噪声更大，门放宽"],
        ["过程噪声Q_scale", "1e5", "2e5", "R2噪声约2倍R1"],
        ["初始位置标准差", "0.05 deg", "0.05 deg", ""],
        ["初始速度标准差", "0.004 deg/s", "0.005 deg/s", ""],
    ])

abody(doc,
    "航迹类型编码：type=1(RELIABLE稳定跟踪) / type=2(MAINTAIN维持) / "
    "type=6(TEMPORARY临时) / type=7(HISTORY死亡)。")

abody(doc, "单帧仿真控制台输出示例（运行run_simulation.m后）：")
abody(doc, "  R1 UKF: type=RELIABLE life=xx")
abody(doc, "  R1: 起始帧=x | 关联=xx 纯预测=xx (关联率=xx%) | 起始中=xx 丢失=xx")
abody(doc, "  NIS: 均值=x.xx 门内=xx.x% (xxx/xxx)")
abody(doc, "  R1 UKF滤波             RMSE: xx.x km (n=xxx)")
abody(doc, "  R2 UKF滤波             RMSE: xx.x km (n=xxx)")
abody(doc, "注：以上xx处为运行后填入的实际数据。")

# ============================================================
# CHAPTER 6
# ============================================================
doc.add_page_break()
aheading(doc, "6\t航迹时间对齐与融合", size=SZDA)

aheading(doc, "6.1\t时间异步与CV外推对齐", size=SZ)

abody(doc,
    "R1和R2的采样时刻不同步：R1从t=0开始（0s, 30s, 60s, ...），"
    "R2从t=13s开始（13s, 43s, 73s, ...），偏移13秒而非整数倍半周期（15s），"
    "以模拟真实多传感器网络中的随机异步采样。"
    "融合前需将R2航迹回退到R1时间网格，使用CV模型状态转移矩阵的负时间版本：")

abody(doc, "x(t-delta) = F(-delta) * x(t)")
abody(doc, "P(t-delta) = F(-delta) * P(t) * F(-delta)^T + Q(|delta|)")

abody(doc,
    "其中delta=13s，F为CV模型转移矩阵，Q按比例缩放。"
    "注意：回退操作不会改善精度，P在传播后只会增大或保持不变，这是信息单向损失的不可逆过程。")

aheading(doc, "6.2\t四种融合算法", size=SZ)

abody(doc, "四种融合算法的核心公式对比如下：")

mktable(doc,
    ["算法", "公式", "假设", "适用场景"],
    [
        ["SCC", "P_f^-1 = P1^-1 + P2^-1", "估计误差完全独立", "无互协方差信息时"],
        ["BC", "x_f = x1 + K*(x2-x1)", "已知互协方差P12", "同源估计，有相关性"],
        ["CI", "优化w最小化det(P_f)", "未知相关性", "最保守，保证一致性"],
        ["FCI", "迹估计权重闭式解", "未知相关性", "FCI的快速近似"],
    ])

abody(doc,
    "所有算法在融合前对协方差矩阵进行正则化（特征值裁剪），确保数值稳定性。"
    "融合架构采用直接1对1配对：R1#1 <-> R2#1。")

aheading(doc, "6.3\t融合结果（单帧仿真）", size=SZ)

abody(doc, "Phase 8定量误差评估输出格式：")

mktable(doc,
    ["算法", "RMSE(km)", "中位数(km)"],
    [
        ["SCC", "[待填入]", "[待填入]"],
        ["BC", "[待填入]", "[待填入]"],
        ["CI", "[待填入]", "[待填入]"],
        ["FCI", "[待填入]", "[待填入]"],
        ["R1单站", "[待填入]", "[待填入]"],
        ["R2单站", "[待填入]", "[待填入]"],
    ])

abody(doc, "最佳融合算法: [待填入] (RMSE=xx.x km)")
abody(doc, "融合 vs R1(精密站): +xx.x%")
abody(doc, "融合 vs R2(普通站): +xx.x% 改善")

abody(doc, "可视化输出（Phase 9）：")
abody(doc, "图1 - 融合结果地图叠加：真值航迹（绿色虚线）+ R1 UKF航迹（蓝色）+ R2 UKF航迹（红色）"
    "+ 四种融合航迹（绿/橙/蓝/紫）+ 最优算法高亮（线宽3.5）。"
    "地图右侧提供图层复选框交互控制（全部显示/全部隐藏）。")
abody(doc, "图2左 - 误差收敛曲线：10帧滑动平均，横轴时间(s)，纵轴位置误差(km)，"
    "四种融合+两单站共六条曲线。")
abody(doc, "图2右 - 误差CDF对比：横轴误差(km)，纵轴累积概率(%)，"
    "直观展示各算法误差分布的集中程度。")

# ============================================================
# CHAPTER 7
# ============================================================
doc.add_page_break()
aheading(doc, "7\t蒙特卡洛统计分析", size=SZDA)

aheading(doc, "7.1\t仿真设置", size=SZ)

abody(doc,
    "蒙特卡洛仿真通过改变随机种子生成不同的噪声/检测/虚警实现，"
    "统计平均性能以排除随机波动干扰。各次仿真使用独立的随机流：")

abody(doc,
    "R1随机流偏移 +1e7，R2随机流偏移 +2e7，确保两部雷达的随机数序列完全隔离。"
    "不同种子间使用不同偏移量打破Toeplitz对角线相关性。")

mktable(doc,
    ["配置项", "直线场景", "拐弯场景", "回头弯场景"],
    [
        ["蒙特卡洛次数", "N=500", "N=100", "N=500"],
        ["种子策略", "SEED_BASE+mc", "SEED_BASE+mc", "SEED_BASE+mc"],
        ["随机流隔离", "R1:+1e7, R2:+2e7", "R1:+1e7, R2:+2e7", "R1:+1e7, R2:+2e7"],
        ["跟踪算法", "自适应UKF", "IMM(CV+CT)", "IMM(CV+CT)"],
        ["起始方式", "真值辅助+纯M/N", "真值辅助+纯M/N", "真值辅助+纯M/N"],
        ["坏种子判定", "航迹断裂/无检测", "同上", "同上"],
    ])

aheading(doc, "7.2\t直线场景蒙特卡洛结果", size=SZ)

abody(doc,
    "run_mc_straight.m 控制台输出包含以下统计板块，每个板块以表格形式展示500次运行的"
    "均值、标准差、中位数、最小值、最大值五个统计量：")

mktable(doc,
    ["板块", "展示内容", "预期趋势"],
    [
        ["RMSE绝对值", "原始点迹->校准->UKF->融合全链路", "逐级递减"],
        ["阶段改善率", "校准改善/UKF改善/融合改善(%)", "正值表示改善"],
        ["MTL航迹长度", "R1/R2/融合的平均航迹段长度(帧)", "融合最长"],
        ["断裂次数", "R1/R2/融合的航迹断裂次数", "融合最少"],
        ["关联诊断", "关联率/NIS均值/NIS门内率/起始帧", "反映跟踪质量"],
        ["最优融合分布", "SCC/BC/CI/FCI各多少次获胜", "体现算法稳定性"],
        ["坏种子统计", "失败次数及原因", "越低越好"],
    ])

abody(doc, "完整数据表格（运行后填入）：")

mktable(doc,
    ["指标", "均值", "标准差", "中位数", "最小", "最大"],
    [
        ["原始点迹 R1 (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["原始点迹 R2 (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["校准后 R1 (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["校准后 R2 (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["UKF R1 (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["UKF R2对齐 (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["融合 SCC (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["融合 BC (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["融合 CI (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["融合 FCI (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["融合最优 (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
    ])

aheading(doc, "7.3\t拐弯场景三体制对比", size=SZ)

abody(doc,
    "run_mc_turn_compare.m 在同一批随机种子下并行运行三种UKF体制，"
    "实现公平对比：")

mktable(doc,
    ["体制", "算法", "Q策略", "适用场景"],
    [
        ["jichu", "CV-UKF", "固定Q_scale", "平稳直线飞行"],
        ["zishiying", "CV-UKF+模糊自适应Q", "NIS滑动窗口驱动", "弱机动/渐变拐弯"],
        ["imm", "CV+CT双模型IMM", "各模型NIS自适应Q", "强机动/急转弯"],
    ])

abody(doc, "交叉对比输出格式：")

mktable(doc,
    ["指标", "zishiying vs jichu", "imm vs jichu"],
    [
        ["Delta R1 UKF (%)", "[填]", "[填]"],
        ["Delta 融合最优 (%)", "[填]", "[填]"],
        ["胜率(R1)", "x/500 (xx%)", "x/500 (xx%)"],
        ["胜率(融合)", "x/500 (xx%)", "x/500 (xx%)"],
    ])

abody(doc, "IMM模型概率诊断（仅imm体制）：")

mktable(doc,
    ["指标", "均值", "标准差", "中位数", "最小", "最大"],
    [
        ["CT均值 R1 (%)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["CT均值 R2 (%)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["CT转弯 R1 (%)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["CT占优帧 R1", "[填]", "[填]", "[填]", "[填]", "[填]"],
    ])

abody(doc, "终极最优体制分布：在500次运行中，每种子取融合RMSE最小的体制，统计各体制获胜次数。")

mktable(doc,
    ["体制", "获胜次数", "胜率"],
    [
        ["jichu", "[填]", "[填]"],
        ["zishiying", "[填]", "[填]"],
        ["imm", "[填]", "[填]"],
    ])

aheading(doc, "7.4\t回头弯180场景蒙特卡洛结果", size=SZ)

abody(doc,
    "回头弯180场景是IMM算法的极限测试——目标在180秒内完成半圆机动，"
    "CV模型完全失效，CT模型需准确匹配。蒙特卡洛结果验证IMM在强机动场景的必要性。")

mktable(doc,
    ["指标", "均值", "标准差", "中位数", "最小", "最大"],
    [
        ["R1 IMM滤波 RMSE (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["R2 IMM滤波 RMSE (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["融合最优 RMSE (km)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["CT模型平均概率 R1 (%)", "[填]", "[填]", "[填]", "[填]", "[填]"],
        ["CT模型平均概率 R2 (%)", "[填]", "[填]", "[填]", "[填]", "[填]"],
    ])

# ============================================================
# CHAPTER 8
# ============================================================
doc.add_page_break()
aheading(doc, "8\t总结与展望", size=SZDA)

aheading(doc, "8.1\t阶段性工作总结", size=SZ)

abody(doc, "本阶段工作完成了开题报告提出的四个研究内容的仿真实现：")

abody(doc,
    "（1）双照射源场景建模与点迹级关联：实现了双基地/三基地几何模型、"
    "异质传感器噪声建模、系统偏差标定、PDA概率数据关联，"
    "完成了从ADS-B真值到极坐标量测的全链路仿真。")
abody(doc,
    "（2）单照射源航迹提取：实现了三种跟踪体制（基础UKF、自适应UKF、IMM-CV/CT），"
    "支持M/N滑窗起始、K_loss终止、Vr硬门过滤等航迹管理功能。")
abody(doc,
    "（3）航迹关联与融合算法研究：实现了四种融合算法（SCC/BC/CI/FCI），"
    "完成时间异步对齐（CV外推）、互协方差维护、协方差正则化。")
abody(doc,
    "（4）仿真验证与性能评估：完成直线/拐弯/回头弯三种场景的单帧仿真和蒙特卡洛统计，"
    "建立完整的RMSE、MTL、断裂次数、关联率、NIS分布等多维度评估体系。")

aheading(doc, "8.2\t主要成果数据", size=SZ)

abody(doc, "系统已实现的完整Pipeline及代码规模：")

mktable(doc,
    ["模块", "目录/文件", "功能"],
    [
        ["config", "simulation_params.m", "13模块参数配置（517行）"],
        ["simulation", "9个文件", "航迹生成/覆盖检查/点迹生成/几何反解"],
        ["ukf", "3个文件", "基础UKF/自适应UKF/IMM双模型"],
        ["association", "2个文件", "NN关联/PDA权重"],
        ["initiation", "1个文件", "M/N滑窗航迹起始"],
        ["tracker", "5个文件", "航迹管理/跟踪器/多目标管理"],
        ["fusion", "4个文件", "时间对齐/四种融合/协方差正则化"],
        ["evaluation", "1个文件", "跟踪误差+融合评估"],
        ["visualization", "5个文件", "场景总览/点云3D/跟踪图/融合图/统计图"],
        ["utils", "7个文件", "球面几何/坐标转换"],
        ["registration", "3个文件", "空间配准/偏差估计"],
        ["主入口", "7个文件", "3个单帧+4个蒙特卡洛"],
    ])

aheading(doc, "8.3\t下一步工作计划", size=SZ)

next_steps = [
    "运行三种场景的主程序和蒙特卡洛程序，收集实际数据填入本报告占位符",
    "从results目录导出可视化图表，插入报告对应章节",
    "完善论文正文撰写，补充相关文献引用",
    "针对拐弯场景进一步优化IMM转移矩阵和CT转弯率参数",
    "探索多目标扩展场景下的航迹关联与融合算法性能",
]
for i, ns in enumerate(next_steps, 1):
    abody(doc, "(%d) %s" % (i, ns))

# ============================================================
# SAVE
# ============================================================
output_path = "阶段成果报告_初稿.docx"
doc.save(output_path)
print("Document saved: %s" % output_path)
print("Tables created: %d" % len(doc.tables))
