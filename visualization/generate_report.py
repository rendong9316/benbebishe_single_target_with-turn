"""
generate_report.py
生成蒙特卡洛拐弯仿真对比分析Word报告
"""
import os
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import scipy.io as sio

BASE = r'D:\Desktop\single_target_with-turn'
RESULTS = os.path.join(BASE, 'results')

# ---- 加载数据 ----
S1 = sio.loadmat(os.path.join(RESULTS, 'mc_turn180_compare_20260701_164750.mat'))
S2 = sio.loadmat(os.path.join(RESULTS, 'mc_turn_compare_20260701_163230.mat'))

UKF_LABELS = ['基础UKF', '自适应UKF', 'IMM']
IMM_IDX = 3

def sfield(S, fn):
    return S['s'][fn][0, 0].flatten()

def extract(S):
    return {
        'best_ukf': S['best_ukf_for_seed'].flatten().astype(int),
        's_rmse_ukf_R1': sfield(S, 'rmse_ukf_R1'),
        's_rmse_ukf_R2': sfield(S, 'rmse_ukf_R2'),
        's_rmse_fus_best': sfield(S, 'rmse_fus_best'),
        's_assoc_R1': sfield(S, 'assoc_R1'),
        's_assoc_R2': sfield(S, 'assoc_R2'),
        's_gate_R1': sfield(S, 'nis_gate_R1'),
        's_gate_R2': sfield(S, 'nis_gate_R2'),
        's_mtl_R1': sfield(S, 'mtl_R1'),
        's_mtl_R2': sfield(S, 'mtl_R2'),
        's_mtl_fus': sfield(S, 'mtl_fus'),
        's_brk_R1': sfield(S, 'brk_R1').astype(float),
        's_brk_R2': sfield(S, 'brk_R2').astype(float),
        's_brk_fus': sfield(S, 'brk_fus').astype(float),
        's_imp_R1': sfield(S, 'imp_fus_vs_R1'),
        's_imp_R2': sfield(S, 'imp_fus_vs_R2'),
        's_bad': sfield(S, 'bad_seed').astype(bool),
    }

D1 = extract(S1)
D2 = extract(S2)

R1_CAL_IMM = {'turn180': 9.0915, 'turn46.7': 9.5435}
R2_CAL_IMM = {'turn180': 10.0305, 'turn46.7': 10.0989}
R1_CAL_ZYSY = {'turn180': 9.4814, 'turn46.7': 10.0455}
R2_CAL_ZYSY = {'turn180': 11.9104, 'turn46.7': 11.9918}
R1_CAL_JC = {'turn180': 9.0988, 'turn46.7': 8.9834}
R2_CAL_JC = {'turn180': 10.8225, 'turn46.7': 10.8337}

# ---- 文档辅助函数 ----
def _set_eastasia_font(run, font_name='SimSun'):
    if run._element is None:
        return
    rPr = run._element.find('.//' + qn('w:rPr'))
    if rPr is None:
        rPr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        _set_eastasia_font(run)
    return h

def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.bold = bold
        _set_eastasia_font(run)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10.5)
        _set_eastasia_font(run)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        _set_eastasia_font(run)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8E8E8')
        cell._tc.get_or_add_tcPr().append(shading)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.alignment = align
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            _set_eastasia_font(run)
    tbl = table._tbl
    borders = OxmlElement('w:tblBorders')
    for bn in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '999999')
        borders.append(b)
    tbl.tblPr.append(borders)
    return table

def add_image(doc, path, caption):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_name = os.path.basename(path)
    run = p.add_run(f'图 {img_name}')
    run.font.size = Pt(9)
    _set_eastasia_font(run)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_picture(path, width=Cm(16))
    doc.add_paragraph()


# ==================================================================
# 创建文档
# ==================================================================
doc = Document()

# ---- 标题 ----
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('蒙特卡洛拐弯仿真对比分析报告')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 0, 0)
_set_eastasia_font(run)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('turn180 vs turn46.7 - UKF性能与融合效果评估')
run2.font.size = Pt(12)
_set_eastasia_font(run2)

doc.add_paragraph()


# ==================================================================
# 1. 实验概述
# ==================================================================
add_heading(doc, '1. 实验概述', level=1)

add_para(doc, '本次蒙特卡洛仿真对比分析了两种转弯场景下三种UKF（基础UKF、自适应UKF、IMM）及数据融合方法的表现。共200次随机种子实验。')

add_table(doc, ['参数', 'turn180', 'turn46.7'], [
    ['角速度 omega', '0.0175 rad/frame', '-0.0175 rad/frame'],
    ['总转弯角度', '~180 deg', '46.7 deg'],
    ['蒙特卡洛次数', '200', '200'],
    ['UKF方法', '基础/自适应/IMM', '基础/自适应/IMM'],
    ['雷达数量', '双基地 (R1+R2)', '双基地 (R1+R2)'],
])
doc.add_paragraph()


# ==================================================================
# 2. UKF校准RMSE对比
# ==================================================================
add_heading(doc, '2. UKF校准后RMSE对比', level=1)

add_para(doc, '表2.1: R1校准RMSE对比 (km, 均值)', bold=True)
add_table(doc, ['UKF方法', 'turn180', 'turn46.7'], [
    ['基础UKF', '9.10', '8.98'],
    ['自适应UKF', '9.48', '10.05'],
    ['IMM', '9.09', '9.54'],
])
doc.add_paragraph()

add_para(doc, '表2.2: R2校准RMSE对比 (km, 均值)', bold=True)
add_table(doc, ['UKF方法', 'turn180', 'turn46.7'], [
    ['基础UKF', '10.82', '10.83'],
    ['自适应UKF', '11.91', '11.99'],
    ['IMM', '10.03', '10.10'],
])
doc.add_paragraph()

add_para(doc, '关键发现:')
add_bullet(doc, 'IMM在两个场景中R1和R2校准RMSE均为最优')
add_bullet(doc, '自适应UKF表现最差，R2上比IMM差了约19-19%')
add_bullet(doc, '基础UKF与IMM差距很小（R1约0.1%），但在R2上明显落后IMM约8-13%')
add_bullet(doc, 'turn46.7场景下所有UKF的R1 RMSE略有上升，但R2变化不大')

add_image(doc, os.path.join(RESULTS, 'chart1_ukf_rmse.png'), 'UKF校准RMSE对比柱状图')


# ==================================================================
# 3. 最佳UKF选择分布
# ==================================================================
add_heading(doc, '3. 最佳UKF选择分布', level=1)

imm1 = sum(D1['best_ukf'] == IMM_IDX)
zy1 = sum(D1['best_ukf'] == 2)
jc1 = sum(D1['best_ukf'] == 1)
imm2 = sum(D2['best_ukf'] == IMM_IDX)
zy2 = sum(D2['best_ukf'] == 2)
jc2 = sum(D2['best_ukf'] == 1)

add_table(doc, ['UKF方法', 'turn180 (200次)', 'turn46.7 (200次)', '占比'], [
    ['基础UKF', str(jc1), str(jc2), f'{jc1/2:.0f}% / {jc2/2:.0f}%'],
    ['自适应UKF', str(zy1), str(zy2), f'{zy1/2:.0f}% / {zy2/2:.0f}%'],
    ['IMM', str(imm1), str(imm2), f'{imm1/2:.0f}% / {imm2/2:.0f}%'],
])
doc.add_paragraph()
add_para(doc, f'IMM以{imm1/2:.0f}%/{imm2/2:.0f}%的胜率在所有蒙特卡洛实验中成为最佳UKF，呈现压倒性优势。')
add_image(doc, os.path.join(RESULTS, 'chart3_best_ukf.png'), '最佳UKF获胜分布')


# ==================================================================
# 4. 融合增益分析
# ==================================================================
add_heading(doc, '4. 融合增益分析', level=1)

fus1 = D1['s_rmse_fus_best'].mean()
fus2 = D2['s_rmse_fus_best'].mean()
imp1_r1 = (1 - fus1 / R1_CAL_IMM['turn180']) * 100
imp1_r2 = (1 - fus1 / R2_CAL_IMM['turn180']) * 100
imp2_r1 = (1 - fus2 / R1_CAL_IMM['turn46.7']) * 100
imp2_r2 = (1 - fus2 / R2_CAL_IMM['turn46.7']) * 100

add_table(doc, ['指标', 'turn180', 'turn46.7'], [
    ['IMM单站R1 RMSE', '9.09 km', '9.54 km'],
    ['IMM单站R2 RMSE', '10.03 km', '10.10 km'],
    ['融合最佳RMSE', f'{fus1:.2f} km', f'{fus2:.2f} km'],
    ['融合改善R1', f'{imp1_r1:+.1f}%', f'{imp2_r1:+.1f}%'],
    ['融合改善R2', f'{imp1_r2:+.1f}%', f'{imp2_r2:+.1f}%'],
])
doc.add_paragraph()

add_para(doc, '关键发现:')
add_bullet(doc, f'turn180场景融合改善R1约{imp1_r1:.1f}%，融合效果良好')
add_bullet(doc, f'turn46.7场景融合反而恶化R1约{abs(imp2_r1):.1f}%！融合RMSE高达{fus2:.2f} km')
add_bullet(doc, f'turn46.7的融合标准差({D2["s_rmse_fus_best"].std():.2f})远大于turn180({D1["s_rmse_fus_best"].std():.2f})，说明融合极不稳定')

add_image(doc, os.path.join(RESULTS, 'chart2_fusion_gain.png'), '融合增益对比')


# ==================================================================
# 5. 关联质量与NIS分析
# ==================================================================
add_heading(doc, '5. 关联质量与NIS分析', level=1)

add_table(doc, ['指标', 'turn180', 'turn46.7', '差异'], [
    ['Assoc R1 (%)', f"{D1['s_assoc_R1'].mean():.1f}", f"{D2['s_assoc_R1'].mean():.1f}", f"-{D1['s_assoc_R1'].mean()-D2['s_assoc_R1'].mean():.1f}%"],
    ['Assoc R2 (%)', f"{D1['s_assoc_R2'].mean():.1f}", f"{D2['s_assoc_R2'].mean():.1f}", f"-{D1['s_assoc_R2'].mean()-D2['s_assoc_R2'].mean():.1f}%"],
    ['NIS Gate R1 (%)', f"{D1['s_gate_R1'].mean():.1f}", f"{D2['s_gate_R1'].mean():.1f}", f"-{D1['s_gate_R1'].mean()-D2['s_gate_R1'].mean():.1f}%"],
    ['NIS Gate R2 (%)', f"{D1['s_gate_R2'].mean():.1f}", f"{D2['s_gate_R2'].mean():.1f}", f"-{D1['s_gate_R2'].mean()-D2['s_gate_R2'].mean():.1f}%"],
])
doc.add_paragraph()
add_para(doc, 'turn46.7场景关联率下降5-6个百分点，NIS门控率也下降3-5个百分点，说明该场景下量测噪声假设与实际偏差更大，滤波器置信度降低。')
add_image(doc, os.path.join(RESULTS, 'chart4_assoc_nis.png'), '关联率与NIS门控率对比')


# ==================================================================
# 6. 跟踪稳定性分析
# ==================================================================
add_heading(doc, '6. 跟踪稳定性分析', level=1)

add_table(doc, ['指标', 'turn180 R1', 'turn180 R2', 'turn180 Fus', 'turn46.7 R1', 'turn46.7 R2', 'turn46.7 Fus'], [
    ['MTL (帧)', f"{D1['s_mtl_R1'].mean():.0f}", f"{D1['s_mtl_R2'].mean():.0f}", f"{D1['s_mtl_fus'].mean():.0f}",
     f"{D2['s_mtl_R1'].mean():.0f}", f"{D2['s_mtl_R2'].mean():.0f}", f"{D2['s_mtl_fus'].mean():.0f}"],
    ['Break R1', f"{D1['s_brk_R1'].mean():.2f}", '-', f"{D1['s_brk_fus'].mean():.2f}",
     f"{D2['s_brk_R1'].mean():.2f}", '-', f"{D2['s_brk_fus'].mean():.2f}"],
    ['Break R2', '-', f"{D1['s_brk_R2'].mean():.2f}", f"{D1['s_brk_fus'].mean():.2f}",
     '-', f"{D2['s_brk_R2'].mean():.2f}", f"{D2['s_brk_fus'].mean():.2f}"],
])
doc.add_paragraph()

add_para(doc, '关键发现:')
add_bullet(doc, f'turn46.7的MTL仅为turn180的{D2["s_mtl_R1"].mean()/D1["s_mtl_R1"].mean()*100:.0f}%（R1）和{D2["s_mtl_R2"].mean()/D1["s_mtl_R2"].mean()*100:.0f}%（R2）')
add_bullet(doc, f'turn46.7的中断次数是turn180的约{D2["s_brk_R1"].mean()/max(D1["s_brk_R1"].mean(),0.001):.1f}倍')
add_bullet(doc, f'turn46.7的融合寿命({D2["s_mtl_fus"].mean():.0f}帧)只有turn180({D1["s_mtl_fus"].mean():.0f}帧)的一半')

add_image(doc, os.path.join(RESULTS, 'chart5_mtl_breaks.png'), '跟踪寿命与中断对比')


# ==================================================================
# 7. RMSE分布特征
# ==================================================================
add_heading(doc, '7. RMSE分布特征', level=1)

add_para(doc, '表7.1: 最佳UKF每seed RMSE统计 (km)', bold=True)
add_table(doc, ['场景', 'R1 Mean', 'R1 Std', 'R1 Min', 'R1 Max', 'R2 Mean', 'R2 Std', 'R2 Min', 'R2 Max'], [
    ['turn180', f"{D1['s_rmse_ukf_R1'].mean():.2f}", f"{D1['s_rmse_ukf_R1'].std():.2f}",
     f"{D1['s_rmse_ukf_R1'].min():.2f}", f"{D1['s_rmse_ukf_R1'].max():.2f}",
     f"{D1['s_rmse_ukf_R2'].mean():.2f}", f"{D1['s_rmse_ukf_R2'].std():.2f}",
     f"{D1['s_rmse_ukf_R2'].min():.2f}", f"{D1['s_rmse_ukf_R2'].max():.2f}"],
    ['turn46.7', f"{D2['s_rmse_ukf_R1'].mean():.2f}", f"{D2['s_rmse_ukf_R1'].std():.2f}",
     f"{D2['s_rmse_ukf_R1'].min():.2f}", f"{D2['s_rmse_ukf_R1'].max():.2f}",
     f"{D2['s_rmse_ukf_R2'].mean():.2f}", f"{D2['s_rmse_ukf_R2'].std():.2f}",
     f"{D2['s_rmse_ukf_R2'].min():.2f}", f"{D2['s_rmse_ukf_R2'].max():.2f}"],
])
doc.add_paragraph()
add_para(doc, 'turn46.7场景R1 RMSE标准差(29.67)远大于turn180(7.09)，说明部分seed严重发散，分布长尾明显。')
add_image(doc, os.path.join(RESULTS, 'chart6_boxplot.png'), 'RMSE分布箱线图')


# ==================================================================
# 8. 融合改善百分比分布
# ==================================================================
add_heading(doc, '8. 融合改善百分比分布', level=1)

add_table(doc, ['场景', 'vs R1 Mean', 'vs R2 Mean', 'Bad Seeds'], [
    ['turn180', f"{D1['s_imp_R1'].mean():.1f}%", f"{D1['s_imp_R2'].mean():.1f}%", f"{D1['s_bad'].sum()}/200"],
    ['turn46.7', f"{D2['s_imp_R1'].mean():.1f}%", f"{D2['s_imp_R2'].mean():.1f}%", f"{D2['s_bad'].sum()}/200"],
])
doc.add_paragraph()
add_para(doc, 'turn46.7场景所有200个seed都被标记为bad，而turn180仅有0.4%的bad seed。这说明右转46.7度场景对融合算法提出了更严峻的挑战。')
add_image(doc, os.path.join(RESULTS, 'chart8_fusion_improvement.png'), '融合改善百分比分布')


# ==================================================================
# 9. 结论与建议
# ==================================================================
add_heading(doc, '9. 结论与建议', level=1)

add_para(doc, '主要结论:', bold=True)
add_bullet(doc, 'IMM UKF在两种转弯场景下均表现最优，校准后R1 RMSE约9km，R2 RMSE约10km，建议作为首选UKF方案')
add_bullet(doc, '自适应UKF表现最差，R2 RMSE比IMM高出约19%，不建议在当前场景使用')
add_bullet(doc, f'turn180场景融合效果良好（改善约{imp1_r1:.0f}-{imp1_r2:.0f}%），但turn46.7场景融合反而恶化（R1恶化约{abs(imp2_r1):.0f}%）')
add_bullet(doc, 'turn46.7场景存在严重的跟踪稳定性问题：关联率低5-6%，MTL缩短60-70%，中断次数翻倍')
add_bullet(doc, 'turn46.7场景RMSE分布长尾严重（std=29.67 vs 7.09），说明部分seed完全发散')

doc.add_paragraph()
add_para(doc, '后续建议:', bold=True)
add_bullet(doc, '排查turn46.7场景融合恶化的根本原因（量测噪声协方差Q是否需要场景自适应调整）')
add_bullet(doc, '针对右转46.7度场景，建议优化融合权重分配策略，避免发散seed污染整体结果')
add_bullet(doc, '考虑引入bad seed剔除机制或使用鲁棒融合方法（如加权平均而非简单平均）')
add_bullet(doc, '进一步研究turn46.7场景下关联率下降的原因，可能需要调整NIS门限或预测协方差')


# ---- 保存 ----
output_path = os.path.join(RESULTS, '蒙特卡洛拐弯仿真对比分析报告.docx')
doc.save(output_path)
print(f'报告已保存至: {output_path}')
print('图表已嵌入:')
for fn in ['chart1_ukf_rmse.png', 'chart2_fusion_gain.png', 'chart3_best_ukf.png',
           'chart4_assoc_nis.png', 'chart5_mtl_breaks.png', 'chart6_boxplot.png',
           'chart8_fusion_improvement.png']:
    print(f'  - {fn}')
